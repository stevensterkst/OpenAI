"""Stable SS Second Brain runtime for v0.8.4.
Read-only by default: no file deletion/rename/move/overwrite and no chat deletion.
"""
from fastapi import FastAPI
from fastapi.responses import HTMLResponse, JSONResponse
from pathlib import Path
from datetime import datetime, timezone
import hashlib, json, mimetypes, os, platform, re, time
import httpx
import psutil
try:
    import keyring
except Exception:
    keyring = None
try:
    from pypdf import PdfReader
except Exception:
    PdfReader = None
try:
    from docx import Document
except Exception:
    Document = None
try:
    from openpyxl import load_workbook
except Exception:
    load_workbook = None

VERSION = "0.8.4"
APP = FastAPI(title="SS Second Brain", version=VERSION)
ROOT = Path(__file__).parent
SERVICE = "SS-Second-Brain"

PROVIDERS = {
    "ollama": {"name":"Ollama", "kind":"local", "base":"http://127.0.0.1:11434", "key":False},
    "jan": {"name":"Jan", "kind":"local", "base":"http://127.0.0.1:1337/v1", "key":False},
    "lmstudio": {"name":"LM Studio", "kind":"local", "base":"http://127.0.0.1:1234/v1", "key":False},
    "openrouter": {"name":"OpenRouter", "kind":"cloud", "base":"https://openrouter.ai/api/v1", "key":True},
    "huggingface": {"name":"Hugging Face", "kind":"cloud", "base":"https://router.huggingface.co/v1", "key":True},
    "venice": {"name":"Venice AI", "kind":"cloud", "base":"https://api.venice.ai/api/v1", "key":True},
    "openai": {"name":"OpenAI", "kind":"cloud", "base":"https://api.openai.com/v1", "key":True},
    "anthropic": {"name":"Anthropic / Claude", "kind":"cloud", "base":"https://api.anthropic.com/v1", "key":True},
    "google": {"name":"Google Gemini", "kind":"cloud", "base":"https://generativelanguage.googleapis.com/v1beta/openai", "key":True},
    "xai": {"name":"xAI / Grok", "kind":"cloud", "base":"https://api.x.ai/v1", "key":True},
    "deepseek": {"name":"DeepSeek", "kind":"cloud", "base":"https://api.deepseek.com/v1", "key":True},
    "mistral": {"name":"Mistral", "kind":"cloud", "base":"https://api.mistral.ai/v1", "key":True},
    "moonshot": {"name":"Kimi / Moonshot", "kind":"cloud", "base":"https://api.moonshot.ai/v1", "key":True},
    "zai": {"name":"Z.ai / GLM", "kind":"cloud", "base":"https://api.z.ai/api/paas/v4", "key":True},
    "qwen": {"name":"Qwen / Alibaba", "kind":"cloud", "base":"https://dashscope-us.aliyuncs.com/compatible-mode/v1", "key":True},
    "perplexity": {"name":"Perplexity", "kind":"cloud", "base":"https://api.perplexity.ai/v1", "key":True},
}
CONNECTORS = {
    "brave": {"name":"Brave Search", "kind":"search", "url":"https://api.search.brave.com/res/v1/web/search", "key":True},
    "duckduckgo": {"name":"DuckDuckGo", "kind":"web", "url":"https://duckduckgo.com", "key":False},
    "tor": {"name":"Tor", "kind":"privacy", "url":"https://www.torproject.org/", "key":False},
    "huggingchat": {"name":"HuggingChat", "kind":"web", "url":"https://huggingface.co/chat/", "key":False},
    "metaai": {"name":"Meta AI", "kind":"web", "url":"https://www.meta.ai/", "key":False},
    "higgsfield": {"name":"Higgsfield", "kind":"web", "url":"https://higgsfield.ai/", "key":False},
}


def data_root():
    root = Path(os.environ.get("LOCALAPPDATA", Path.home())) / "SS" / "data" if os.name == "nt" else Path.home()/".local/share/SS/data"
    for sub in ("chats", "memory", "backups"):
        (root/sub).mkdir(parents=True, exist_ok=True)
    return root


def cloud_root():
    raw = os.environ.get("SS_CHAT_CLOUD_ROOT", "").strip()
    if not raw: return None
    p = Path(raw).expanduser(); p.mkdir(parents=True, exist_ok=True); return p


def safe(v):
    return (re.sub(r"[^A-Za-z0-9._-]+", "_", str(v or "chat"))[:140] or "chat")


def atomic(path, obj):
    path.parent.mkdir(parents=True, exist_ok=True)
    tmp = path.with_suffix(path.suffix + ".tmp")
    tmp.write_text(json.dumps(obj, ensure_ascii=False, indent=2), encoding="utf-8")
    os.replace(tmp, path)


def save_chat(chat):
    now = datetime.now(timezone.utc).isoformat()
    c = dict(chat); c.setdefault("created_at", now); c["updated_at"] = now
    targets = [data_root()/"chats"/(safe(c["id"])+".json")]
    cr = cloud_root()
    if cr: targets.append(cr/(safe(c["id"])+".json"))
    for p in targets: atomic(p, c)
    return {"local":str(targets[0]), "cloud":str(targets[1]) if len(targets)>1 else None}


def chats():
    out=[]
    for p in sorted((data_root()/"chats").glob("*.json"), key=lambda x:x.stat().st_mtime, reverse=True):
        try:
            c=json.loads(p.read_text(encoding="utf-8")); out.append({"id":c.get("id",p.stem),"title":c.get("title","Untitled"),"updated_at":c.get("updated_at"),"provider":c.get("provider"),"model":c.get("model"),"messages":len(c.get("messages",[]))})
        except Exception: pass
    return out


def credential(pid, supplied=None):
    if supplied: return supplied
    if keyring:
        try: return keyring.get_password(SERVICE, pid)
        except Exception: return None
    return None


def boundary(pid, exc):
    s=str(exc); name=PROVIDERS.get(pid, {"name":pid})["name"]
    low=s.lower()
    if "connection refused" in low or "connecterror" in low or "timeout" in low:
        return f"{name} is unreachable. SS stopped at that boundary."
    if "401" in s or "403" in s: return f"{name} rejected the credential or permission."
    if "429" in s: return f"{name} rate-limited the request. SS did not silently spend elsewhere."
    return s


async def request(url, method="GET", headers=None, payload=None, params=None):
    async with httpx.AsyncClient(timeout=120, follow_redirects=True) as client:
        r=await client.request(method, url, headers=headers, json=payload, params=params)
        try: data=r.json()
        except Exception: data={"raw":r.text}
        if r.status_code >= 400:
            raise RuntimeError(f"HTTP {r.status_code}: {data.get('error') if isinstance(data,dict) else r.text}")
        return data


def local_candidates():
    try: free=psutil.virtual_memory().available/2**30
    except Exception: free=0
    # Prefer the strongest model that is plausibly safe for the current RAM budget.
    if free >= 7: return ["qwen3:4b","phi4-mini:3.8b","qwen3:1.7b","gemma3:1b","llama3.2:1b"]
    if free >= 4: return ["phi4-mini:3.8b","qwen3:1.7b","gemma3:1b","llama3.2:1b"]
    if free >= 2: return ["qwen3:1.7b","gemma3:1b","llama3.2:1b"]
    return ["gemma3:1b","llama3.2:1b"]


@APP.get("/", response_class=HTMLResponse)
async def home(): return HTMLResponse((ROOT/"web"/"stable.html").read_text(encoding="utf-8"))

@APP.get("/console", response_class=HTMLResponse)
async def console(): return HTMLResponse((ROOT/"web"/"stable.html").read_text(encoding="utf-8"))

@APP.get("/system")
async def system(): return {"service":"SS Second Brain","version":VERSION,"status":"online","port":8765,"entry":"http://127.0.0.1:8765/","policy":{"auto_delete_chats":False,"cloud_spend_without_request_approval":False,"file_mutation":False},"storage":{"local":str(data_root()),"cloud":str(cloud_root()) if cloud_root() else None}}

@APP.get("/api/providers")
async def provider_list(): return {"providers":PROVIDERS,"connectors":CONNECTORS,"version":VERSION}

@APP.get("/api/resources")
async def resources():
    v=psutil.virtual_memory(); s=psutil.swap_memory()
    return {"ok":True,"ram_total_gb":round(v.total/2**30,2),"ram_available_gb":round(v.available/2**30,2),"ram_used_percent":v.percent,"swap_used_gb":round(s.used/2**30,2),"cpu_percent":psutil.cpu_percent(interval=.05)}

@APP.get("/api/storage")
async def storage(): return {"local":str(data_root()),"cloud":str(cloud_root()) if cloud_root() else None,"never_delete":True}

@APP.get("/api/chats")
async def chat_list(): return {"chats":chats(),"never_delete":True}

@APP.get("/api/chats/{cid}")
async def chat_get(cid):
    p=data_root()/"chats"/(safe(cid)+".json")
    if not p.exists(): return JSONResponse({"ok":False,"error":"Chat not found"},404)
    return json.loads(p.read_text(encoding="utf-8"))

@APP.post("/api/chats")
async def chat_store(body:dict):
    if not body.get("id") or not isinstance(body.get("messages"),list): return JSONResponse({"ok":False,"error":"id and messages required"},400)
    return {"ok":True,"saved":save_chat(body),"never_delete":True}

@APP.get("/api/credentials/status")
async def credential_status():
    ids=list(PROVIDERS)+["brave"]
    return {"available":bool(keyring),"backend":"OS credential store" if keyring else None,"configured":{p:bool(credential(p)) for p in ids}}

@APP.post("/api/credentials")
async def credential_save(body:dict):
    pid=body.get("provider"); key=body.get("apiKey")
    if pid not in PROVIDERS and pid!="brave": return JSONResponse({"ok":False,"error":"Unknown provider"},400)
    if not keyring: return JSONResponse({"ok":False,"error":"OS credential store unavailable"},500)
    if not key: return JSONResponse({"ok":False,"error":"Credential required"},400)
    keyring.set_password(SERVICE,pid,key); return {"ok":True,"provider":pid,"stored":"OS credential store"}

@APP.post("/api/models")
async def models(body:dict):
    pid=body.get("provider"); p=PROVIDERS.get(pid)
    if not p: return JSONResponse({"ok":False,"error":"Unknown provider"},400)
    key=credential(pid,body.get("apiKey")); base=(body.get("base") or p["base"]).rstrip("/")
    if p["key"] and not key: return JSONResponse({"ok":False,"error":f"{p['name']} needs an API key"},400)
    try:
        start=time.perf_counter()
        if pid=="ollama":
            d=await request(base+"/api/tags"); ms=[{"id":m.get("name"),"is_free":True,"detail":m.get("details",{}).get("parameter_size","")} for m in d.get("models",[])]
        else:
            h={"Authorization":f"Bearer {key}"}; d=await request(base+"/models",headers=h); ms=[]
            for m in d.get("data",[]):
                if m.get("id"): ms.append({"id":m["id"],"detail":m.get("owned_by","")})
        ms=sorted(ms,key=lambda x:x["id"].lower()); return {"ok":True,"models":ms,"latency_ms":round((time.perf_counter()-start)*1000)}
    except Exception as e: return JSONResponse({"ok":False,"error":boundary(pid,e)},502)

@APP.post("/api/chat")
async def chat(body:dict):
    pid=body.get("provider"); p=PROVIDERS.get(pid)
    if not p: return JSONResponse({"ok":False,"error":"Unknown provider"},400)
    if p["kind"]=="cloud" and not body.get("cloud_approved"): return JSONResponse({"ok":False,"error":"Cloud request not approved for this turn."},403)
    key=credential(pid,body.get("apiKey")); model=body.get("model"); messages=body.get("messages") or []
    if not model: return JSONResponse({"ok":False,"error":"Model required"},400)
    if p["key"] and not key: return JSONResponse({"ok":False,"error":f"{p['name']} needs an API key"},400)
    try:
        start=time.perf_counter(); base=p["base"].rstrip("/")
        if pid=="ollama":
            payload={"model":model,"messages":messages,"stream":False,"options":{"temperature":body.get("temperature",0.7)}}
            d=await request(base+"/api/chat","POST",payload=payload); text=(d.get("message") or {}).get("content","")
        elif pid=="anthropic":
            h={"x-api-key":key,"anthropic-version":"2023-06-01"}; d=await request(base+"/messages","POST",headers=h,payload={"model":model,"max_tokens":2048,"messages":messages}); text="".join(x.get("text","") for x in d.get("content",[]) if isinstance(x,dict))
        else:
            h={"Authorization":f"Bearer {key}"}; payload={"model":model,"messages":messages,"temperature":body.get("temperature",0.7),"stream":False}
            if pid=="openrouter" and body.get("zdr"): payload["provider"]={"data_collection":"deny"}
            d=await request(base+"/chat/completions","POST",headers=h,payload=payload); text=((d.get("choices") or [{}])[0].get("message") or {}).get("content","")
        result={"ok":True,"text":text,"provider":pid,"model":model,"latency_ms":round((time.perf_counter()-start)*1000)}
        cid=body.get("chat_id")
        if cid: save_chat({"id":cid,"title":body.get("title") or "SS chat","provider":pid,"model":model,"messages":list(messages)+[{"role":"assistant","content":text}]})
        return result
    except Exception as e: return JSONResponse({"ok":False,"error":boundary(pid,e)},502)

@APP.post("/api/auto-chat")
async def auto_chat(body:dict):
    task=(body.get("task") or "").strip(); messages=body.get("messages") or [{"role":"user","content":task}]
    if not task: return JSONResponse({"ok":False,"error":"Task required"},400)
    # First try local models. Cloud is never entered without explicit approval.
    d=await models({"provider":"ollama"})
    if d.get("ok") and d.get("models"):
        available={m["id"] for m in d["models"]}; chosen=next((m for m in local_candidates() if m in available),d["models"][0]["id"])
        return await chat({"provider":"ollama","model":chosen,"messages":messages,"chat_id":body.get("chat_id"),"title":task[:80]})
    return JSONResponse({"ok":False,"error":"No working local model. Start Ollama or explicitly enable a configured cloud provider."},502)

# ---------------- READ-ONLY file intelligence ----------------
TEXT_EXT={".txt",".md",".csv",".json",".xml",".html",".htm",".rtf",".log",".yaml",".yml"}
IMAGE_EXT={".jpg",".jpeg",".png",".webp",".gif",".bmp",".tif",".tiff",".heic"}

def iter_files(folder): return (p for p in Path(folder).expanduser().rglob("*") if p.is_file())
def digest(p):
    h=hashlib.sha256()
    with open(p,"rb") as f:
        for chunk in iter(lambda:f.read(1024*1024),b""): h.update(chunk)
    return h.hexdigest()
def meta(p,hash_it=False):
    s=p.stat(); r={"path":str(p),"name":p.name,"size_bytes":s.st_size,"extension":p.suffix.lower(),"mime":mimetypes.guess_type(p.name)[0],"created_at":datetime.fromtimestamp(s.st_ctime).isoformat(),"modified_at":datetime.fromtimestamp(s.st_mtime).isoformat()}
    if hash_it:r["sha256"]=digest(p)
    return r
def extract(p):
    ext=p.suffix.lower()
    if ext in TEXT_EXT:return p.read_text(encoding="utf-8",errors="replace")
    if ext==".pdf" and PdfReader:return "\n\n".join(x.extract_text() or "" for x in PdfReader(str(p)).pages)
    if ext==".docx" and Document:
        d=Document(str(p)); return "\n".join([x.text for x in d.paragraphs]+[" | ".join(c.text for c in row.cells) for t in d.tables for row in t.rows])
    if ext in (".xlsx",".xlsm") and load_workbook:
        wb=load_workbook(str(p),read_only=True,data_only=True); out=[]
        for ws in wb.worksheets:
            out.append("## "+ws.title); out += [" | ".join("" if v is None else str(v) for v in row) for row in ws.iter_rows(values_only=True)]
        return "\n".join(out)
    raise ValueError(f"No safe extractor for {ext}")

@APP.post("/api/files/scan")
async def file_scan(body:dict):
    root=Path(body.get("folder","")).expanduser()
    if not root.is_dir(): return JSONResponse({"ok":False,"error":f"Folder does not exist: {root}"},400)
    limit=min(int(body.get("limit",5000)),20000); rows=[]
    for i,p in enumerate(iter_files(root)):
        if i>=limit:break
        try: rows.append(meta(p,bool(body.get("hash"))))
        except OSError: pass
    return {"ok":True,"files":rows,"count":len(rows),"truncated":len(rows)>=limit,"read_only":True}

@APP.post("/api/files/extract")
async def file_extract(body:dict):
    p=Path(body.get("path","")).expanduser()
    try:
        text=extract(p); lim=min(int(body.get("max_chars",200000)),500000); return {"ok":True,"metadata":meta(p),"text":text[:lim],"truncated":len(text)>lim,"read_only":True}
    except Exception as e:return JSONResponse({"ok":False,"error":str(e)},400)

@APP.post("/api/files/duplicates")
async def duplicates(body:dict):
    root=Path(body.get("folder","")).expanduser()
    if not root.is_dir(): return JSONResponse({"ok":False,"error":f"Folder does not exist: {root}"},400)
    groups={}; count=0
    for p in iter_files(root):
        if p.suffix.lower() not in IMAGE_EXT:continue
        try:
            m=meta(p,True);groups.setdefault(m["sha256"],[]).append(m);count+=1
        except OSError:pass
    proposals=[]
    for g in groups.values():
        if len(g)<2:continue
        keep=min(g,key=lambda x:(x["created_at"],x["path"].lower())); dup=[x for x in g if x["path"]!=keep["path"]]
        proposals.append({"keep":keep,"duplicates":dup,"recoverable_bytes":sum(x["size_bytes"] for x in dup),"reason":"Exact SHA-256 duplicate; oldest creation timestamp proposed as canonical. No action performed."})
    return {"ok":True,"images_scanned":count,"duplicate_groups":proposals,"recoverable_bytes":sum(x["recoverable_bytes"] for x in proposals),"actions_performed":[],"read_only":True}

@APP.get("/api/files/policy")
async def file_policy(): return {"read_only":True,"delete":False,"rename":False,"move":False,"overwrite":False,"preserve_metadata":True,"permission_required_for_mutation":True}
